"""
Step 4: Cluster Interpretation & Archetype Naming
===================================================
- Detailed cluster profiles (clinical, prescribing, prescriber)
- Top complaints and drugs per cluster
- Suggest archetype names
- Generate Word report
"""
import pandas as pd
import numpy as np
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import ENCOUNTER_CLUSTERED_PATH, REPORTS_DIR, FIGURES_DIR, AGE_STRATA


def analyse_cluster(df, cluster_id):
    """Analyse a single cluster and return a summary dict."""
    c = df[df["cluster"] == cluster_id]
    n = len(c)
    total = len(df)

    summary = {
        "cluster": cluster_id,
        "n": n,
        "pct_of_total": round(n / total * 100, 1),
    }

    # Clinical profile
    summary["mean_age_months"] = round(c["age_months"].mean(), 1)
    summary["mean_weight"] = round(c["PatientWeight"].mean(), 1) if c["PatientWeight"].notna().any() else "N/A"
    summary["dominant_age_group"] = c["age_stratum"].mode().iloc[0] if len(c) > 0 else "N/A"
    summary["gender_dist"] = c["Gender"].value_counts().to_dict()
    summary["visit_type_dist"] = c["VisitType"].value_counts().head(5).to_dict()

    # Prescribing pattern
    summary["ab_rate"] = round(c["has_antibiotic"].mean() * 100, 1)
    summary["mono_rate"] = round(c["antibiotic_monotherapy"].mean() * 100, 1)
    summary["combo_rate"] = round(c["antibiotic_combination"].mean() * 100, 1)
    summary["polypharmacy_rate"] = round(c["polypharmacy_flag"].mean() * 100, 1)
    summary["mean_drugs"] = round(c["num_distinct_drugs"].mean(), 2)
    summary["mean_antibiotics"] = round(c["num_antibiotics"].mean(), 2)

    # Prescriber characteristics
    if "prescriber_antibiotic_rate" in c.columns:
        summary["mean_prescriber_ab_rate"] = round(c["prescriber_antibiotic_rate"].mean() * 100, 1)
        summary["std_prescriber_ab_rate"] = round(c["prescriber_antibiotic_rate"].std() * 100, 1)

    # Top 10 complaints
    complaints = []
    for comp_str in c["Complaint"].dropna():
        for comp in str(comp_str).split(","):
            comp = comp.strip()
            if comp and comp.lower() not in ("nan", ""):
                complaints.append(comp)
    summary["top_complaints"] = Counter(complaints).most_common(10)

    # Top 10 drugs
    drugs = []
    for drug_str in c["drug_names"].dropna():
        for d in str(drug_str).split("|"):
            d = d.strip()
            if d and d.lower() not in ("nan", ""):
                drugs.append(d)
    summary["top_drugs"] = Counter(drugs).most_common(10)

    return summary


def suggest_archetype_name(summary):
    """Suggest a clinically meaningful archetype name based on profile."""
    ab = summary["ab_rate"]
    poly = summary["polypharmacy_rate"]
    combo = summary["combo_rate"]
    age = summary["dominant_age_group"]

    parts = []

    # Antibiotic intensity
    if ab < 20:
        parts.append("Low-Antibiotic")
    elif ab < 50:
        parts.append("Moderate-Antibiotic")
    else:
        parts.append("High-Antibiotic")

    # Polypharmacy
    if poly > 50:
        parts.append("High-Polypharmacy")

    # Combination
    if combo > 30:
        parts.append("Combination-Therapy")
    elif ab > 0 and combo < 10:
        parts.append("Monotherapy-Dominant")

    # Age
    age_short = {
        "neonate_0_27d": "Neonatal",
        "infant_1_2m": "Young-Infant",
        "infant_3_5m": "Infant",
        "infant_6_11m": "Infant",
        "toddler_12_23m": "Toddler",
        "preschool_2_4y": "Preschool",
        "child_5_11y": "School-Age",
        "adolescent_12_17y": "Adolescent",
        "adult_18plus": "Adult",
    }
    parts.append(age_short.get(age, "Mixed-Age"))

    return " / ".join(parts)


def create_word_report(cluster_summaries, df):
    """Generate a professional Word document report."""
    doc = Document()

    # Title
    title = doc.add_heading("Cluster Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Antibiotic Prescribing Archetypes - Paediatric OPD, January 2026\n"
        "Lady Ridgeway Hospital / PGIM University of Colombo"
    )
    doc.add_paragraph("")

    # Executive summary table
    doc.add_heading("Executive Summary", level=1)
    n_clusters = len(cluster_summaries)
    table = doc.add_table(rows=n_clusters + 1, cols=7)
    table.style = "Light Shading Accent 1"
    headers = ["Cluster", "N (%)", "AB Rate%", "Mono%", "Combo%", "Polypharm%", "Archetype"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for idx, s in enumerate(cluster_summaries):
        row = table.rows[idx + 1]
        row.cells[0].text = str(s["cluster"])
        row.cells[1].text = f"{s['n']} ({s['pct_of_total']}%)"
        row.cells[2].text = f"{s['ab_rate']}"
        row.cells[3].text = f"{s['mono_rate']}"
        row.cells[4].text = f"{s['combo_rate']}"
        row.cells[5].text = f"{s['polypharmacy_rate']}"
        row.cells[6].text = s.get("archetype_name", "")

    doc.add_paragraph("")

    # Detailed profiles
    doc.add_heading("Detailed Cluster Profiles", level=1)
    for s in cluster_summaries:
        doc.add_heading(
            f"Cluster {s['cluster']}: {s.get('archetype_name', 'TBD')}",
            level=2
        )
        doc.add_paragraph(
            f"Size: {s['n']} encounters ({s['pct_of_total']}% of total)"
        )

        doc.add_heading("Clinical Profile", level=3)
        doc.add_paragraph(f"Dominant age group: {s['dominant_age_group']}")
        doc.add_paragraph(f"Mean age: {s['mean_age_months']} months")
        doc.add_paragraph(f"Mean weight: {s['mean_weight']} kg")
        doc.add_paragraph(f"Gender: {s['gender_dist']}")

        doc.add_heading("Prescribing Pattern", level=3)
        doc.add_paragraph(f"Antibiotic rate: {s['ab_rate']}%")
        doc.add_paragraph(f"Monotherapy: {s['mono_rate']}%")
        doc.add_paragraph(f"Combination therapy: {s['combo_rate']}%")
        doc.add_paragraph(f"Polypharmacy: {s['polypharmacy_rate']}%")
        doc.add_paragraph(f"Mean drugs/encounter: {s['mean_drugs']}")
        doc.add_paragraph(f"Mean antibiotics/encounter: {s['mean_antibiotics']}")

        if "mean_prescriber_ab_rate" in s:
            doc.add_heading("Prescriber Characteristics", level=3)
            doc.add_paragraph(
                f"Mean prescriber AB rate: {s['mean_prescriber_ab_rate']}% "
                f"(SD: {s['std_prescriber_ab_rate']}%)"
            )

        doc.add_heading("Top 10 Complaints", level=3)
        for complaint, count in s["top_complaints"]:
            doc.add_paragraph(f"{complaint}: {count}", style="List Bullet")

        doc.add_heading("Top 10 Drugs", level=3)
        for drug, count in s["top_drugs"]:
            doc.add_paragraph(f"{drug}: {count}", style="List Bullet")

        doc.add_paragraph("")

    # Add figures if they exist
    import os
    doc.add_heading("Visualizations", level=1)
    figure_files = [
        "10_pca_clusters.png",
        "11_cluster_heatmap.png",
        "12_cluster_comparison_bars.png",
        "13_cluster_sizes.png",
    ]
    for fig_file in figure_files:
        fig_path = os.path.join(FIGURES_DIR, fig_file)
        if os.path.exists(fig_path):
            doc.add_picture(fig_path, width=Inches(5.5))
            doc.add_paragraph("")

    report_path = f"{REPORTS_DIR}/cluster_analysis_report.docx"
    doc.save(report_path)
    print(f"  Report saved: {report_path}")


def run():
    print("=" * 60)
    print("STEP 4: ARCHETYPE INTERPRETATION")
    print("=" * 60)

    df = pd.read_csv(ENCOUNTER_CLUSTERED_PATH, low_memory=False)
    print(f"\nLoaded: {len(df)} encounters with {df['cluster'].nunique()} clusters")

    cluster_ids = sorted(df["cluster"].unique())
    summaries = []

    for cid in cluster_ids:
        print(f"\n--- Cluster {cid} ---")
        s = analyse_cluster(df, cid)
        s["archetype_name"] = suggest_archetype_name(s)
        print(f"  Suggested archetype: {s['archetype_name']}")
        print(f"  N={s['n']} ({s['pct_of_total']}%), AB rate={s['ab_rate']}%, "
              f"Polypharmacy={s['polypharmacy_rate']}%")
        print(f"  Top 3 complaints: {[c[0] for c in s['top_complaints'][:3]]}")
        print(f"  Top 3 drugs: {[d[0] for d in s['top_drugs'][:3]]}")
        summaries.append(s)

    create_word_report(summaries, df)
    return summaries


if __name__ == "__main__":
    run()
